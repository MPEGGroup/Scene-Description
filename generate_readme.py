import os
import re
import markdown
from docx import Document

def convert_table_to_markdown(table):
    markdown_table = ""    
    headers = table.rows[0].cells
    markdown_table += "| " + " | ".join(cell.text.strip() for cell in headers) + " |\n"
    markdown_table += "|-" + "-|-".join(["" for _ in headers]) + "-|\n"
    for row in table.rows[1:]:
        markdown_table += "| " + " | ".join(cell.text.strip() for cell in row.cells) + " |\n"
    return markdown_table

def index_of_table(doc, elem):
    index = 0
    for tbl in doc.tables:
        if tbl._element == elem:
            #print(f"Found table at index {index}")
            return doc.tables[index]
        index += 1

def extract_subsections(doc, section_number):
    subsections = []
    current_subsection = None    
    
    for para in doc.paragraphs:        
        if para.text and para.text[0].isdigit() and len(para.text.split('.')) <= 2:  # Check for 2nd level sections
            if current_subsection:
                subsections.append(current_subsection)
                current_subsection = None
        
        if para.text.startswith(f"{section_number}.") and len(para.text.split('.')) == 2:
            current_subsection = {'title': para.text, 'content': []}
        elif current_subsection:
            current_subsection['content'].append(para.text)
        
            # Check for tables after the current paragraph in the document
            next_elem = para._element.getnext()
            if next_elem is not None and next_elem.tag.endswith('tbl'): 
                markdown_table = convert_table_to_markdown(index_of_table(doc, next_elem))
                current_subsection['content'].append(markdown_table)
                    
    return subsections

def convert_to_markdown(subsections, output_pattern):
    for idx, subsection in enumerate(subsections):
        title = subsection['title']
        fn_title = title.strip().replace(' ', '_').replace('\t','')
        file_name = f"{output_pattern}_{fn_title}.md"
        content = "\n".join(subsection['content'])
        
        markdown_content = f"# {title}\n\n{content}\n"
        
        with open(file_name, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
        
        print(f"Converted {title} to {file_name}")

def main(doc_name, output_pattern):
    doc = Document(doc_name)
    section_number = 8
    subsections = extract_subsections(doc, section_number)
    convert_to_markdown(subsections, output_pattern)

if __name__ == "__main__":    
    main("MDS23831_WG03_N01221_23090-14_2nd_edition_rm.docx", "README")
